from flask import Flask, render_template, request, send_file
import requests
import io
from docx import Document

app = Flask(__name__, template_folder="templates", static_folder="static")
app.config['TEMPLATES_AUTO_RELOAD'] = True

# Function to fetch school district/university data
def get_district_data(name, location):
    search_queries = {
        "Website": f"{name} {location} school district site:.gov OR site:.edu",
        "Key Stakeholders": f"{name} {location} school district leadership site:.gov OR site:.edu",
        "Enrollment": f"{name} {location} school district enrollment number of high schools site:.gov OR site:.edu",
        "Wellness Initiatives": f"{name} {location} school district wellness program site:.gov OR site:.edu",
        "Grants": f"{name} {location} school district available grants wellness funding site:.gov OR site:.edu",
        "News": f"{name} {location} school district news board decisions site:.gov OR site:.edu"
    }

    search_results = {}

    for category, query in search_queries.items():
        search_url = f"https://www.googleapis.com/customsearch/v1?q={query}&key=AIzaSyDkW3m1DdVfHbqirYNRa2eSna1ZSG3BY_o&cx=77cbed915d0c641b3"
        response = requests.get(search_url)

        if response.status_code == 200:
            results = response.json().get("items", [])
            if results:
                search_results[category] = {
                    "title": results[0].get("title", "No Title"),
                    "snippet": results[0].get("snippet", "No description available."),
                    "link": results[0].get("link", "#")
                }
            else:
                search_results[category] = {"title": "No Data Found", "snippet": "", "link": "#"}
        else:
            search_results[category] = {"title": "Error retrieving data", "snippet": "", "link": "#"}

    return search_results

# Function to create Word document
def create_word_doc(data):
    doc = Document()
    doc.add_heading('School District Report', level=1)
    for key, value in data.items():
        doc.add_paragraph(f"{key}: {value}")
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        district_name = request.form.get('district_name', '').strip()
        location = request.form.get('location', '').strip()

        if not district_name or not location:
            return render_template("index.html", error="Please fill out both fields.")

        search_results = get_district_data(district_name, location)

        print("Search Results:", search_results)  # Debugging

        data = {
            "District Name": district_name,
            "Location": location,
            "Results": search_results  # Pass multiple search results
        }

        return render_template("index.html", data=data)

    return render_template("index.html")


@app.route('/export', methods=['POST'])
def export():
    data = {
        "District Name": request.form['district_name'],
        "Location": request.form['location'],
        "Website": request.form['website']
    }
    file_stream = create_word_doc(data)
    return send_file(file_stream, as_attachment=True, download_name="district_report.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route('/test')
def test():
    return "Flask is running correctly!"

@app.route('/debug-templates')
def debug_templates():
    import os
    template_path = os.path.join(app.template_folder, "index.html")
    return f"Flask is looking for: {template_path}. Exists: {os.path.exists(template_path)}"

if __name__ == '__main__':
    app.run(debug=True, host="0.0.0.0", port=10000)
